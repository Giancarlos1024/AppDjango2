from django.shortcuts import render, redirect, reverse
from django.http import HttpResponse
from docx import Document
from docx2pdf import convert
from .forms import SpecialTestimonyForm, IvaSalesTestimonyForm, MunicipalityNoticesForm, BuySellForm
import os
import pythoncom
import base64
from .models import SpecialTestimony, IvaSalesTestimony, MunicipalityNotices, BuySell


# Create your views here.
def dashboard_view(request):
    # Contar el número de registros en cada modelo
    total_special_testimony = SpecialTestimony.objects.count()
    total_iva_sales_testimony = IvaSalesTestimony.objects.count()
    total_municipality_notices = MunicipalityNotices.objects.count()
    total_buy_sell = BuySell.objects.count()

    # Pasar los totales al contexto del template
    context = {
        'total_special_testimony': total_special_testimony,
        'total_iva_sales_testimony': total_iva_sales_testimony,
        'total_municipality_notices': total_municipality_notices,
        'total_buy_sell': total_buy_sell,
    }

    return render(request, 'dashboard.html', context)


# views of options documents
def option_generate_document(request):
    options_documents = ['COMPRAVENTA SIMPLE', 'AVISOS MUNICIPALIDAD DICABI', 'TESTIMONIO COMPRAVENTA IVA',
                         'TESTIMONIO ESPECIAL']
    return render(request, 'documents/option-generate-document.html', {
        'options_documents': options_documents
    })


# Vista para generar el documento (sin descargas)
def replace_text_in_runs(paragraph, replacements):
    """
    Reemplaza el texto en los runs de un párrafo sin perder el formato,
    incluso si el texto a reemplazar está dividido en múltiples runs.
    """
    for key, value in replacements.items():
        if key in paragraph.text:
            # Crear una cadena temporal para almacenar el texto completo del párrafo
            full_text = ''.join([run.text for run in paragraph.runs])
            # Reemplazar la clave por el valor en el texto completo
            new_text = full_text.replace(key, value)

            # Limpiar el texto de los runs
            for run in paragraph.runs:
                run.text = ''

            # Dividir el nuevo texto en partes según los runs originales
            parts = new_text.split(key)
            for i, run in enumerate(paragraph.runs):
                if i < len(parts):
                    run.text = parts[i]
                else:
                    run.text = ''

            # Si hay más partes que runs, agregar nuevos runs
            if len(parts) > len(paragraph.runs):
                for part in parts[len(paragraph.runs):]:
                    new_run = paragraph.add_run(part)
                    new_run.bold = paragraph.runs[-1].bold  # Mantener negrita
                    new_run.italic = paragraph.runs[-1].italic  # Mantener cursiva
                    new_run.font.size = paragraph.runs[-1].font.size  # Mantener tamaño de fuente


def replace_text_in_docx(docx_path, replacements, output_path):
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacements)

    doc.save(output_path)


def generate_special_testimony_document(request):
    pythoncom.CoInitialize()

    try:
        if request.method == 'POST':
            form = SpecialTestimonyForm(request.POST)
            if form.is_valid():
                deed_numbers_letters = form.cleaned_data['deed_numbers_letters']
                deed_numbers_digits = form.cleaned_data['deed_numbers_digits']
                municipality = form.cleaned_data['municipality']
                department = form.cleaned_data['department']
                day_letters = form.cleaned_data['day_letters']
                day_digits = form.cleaned_data['day_digits']
                month_letters_a = form.cleaned_data['month_letters_a']
                month_letters_g = form.cleaned_data['month_letters_g']
                year_letters = form.cleaned_data['year_letters']
                year_digits = form.cleaned_data['year_digits']
                notary_office = form.cleaned_data['notary_office']
                name_notary_office = form.cleaned_data['name_notary_office']
                number_of_pages = form.cleaned_data['number_of_pages']
                she = form.cleaned_data['she']

                if notary_office == 'notario':
                    she = 'el'
                else:
                    she = 'la'

                BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                document_special_testimony_path = os.path.join(BASE_DIRECTORY, 'documents', 'Testimonio_Especial.docx')

                replacements = {
                    "(numero_escritura_letras)": deed_numbers_letters,
                    "(numero_escritura_numeros)": deed_numbers_digits,
                    "(municipio)": municipality,
                    "(departamento)": department,
                    "(dia_letras)": day_letters,
                    "(dia_numeros)": day_digits,
                    "(mes_letras_a)": month_letters_a,
                    "(mes_letras_g)": month_letters_g,
                    "(anio_letras)": year_letters,
                    "(anio_numeros)": year_digits,
                    "(notario_notaria)": notary_office,
                    "(nombre_notario_notaria)": name_notary_office,
                    "(numero_hojas)": number_of_pages,
                    "(el_la)": she
                }

                temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Testimonio_Especial.docx')
                replace_text_in_docx(document_special_testimony_path, replacements, temp_docx_path)

                request.session['word_generated'] = True
                return redirect('download_special_testimony_docx')

        else:
            form = SpecialTestimonyForm()

        return render(request, 'documents/special-testimony-document.html', {'form': form})

    finally:
        pythoncom.CoUninitialize()


def generate_iva_sales_testimony_document(request):
    pythoncom.CoInitialize()

    try:
        if request.method == 'POST':
            form = IvaSalesTestimonyForm(request.POST)
            if form.is_valid():
                deed_numbers_letters = form.cleaned_data['deed_numbers_letters']
                deed_numbers_digits = form.cleaned_data['deed_numbers_digits']
                municipality = form.cleaned_data['municipality']
                department = form.cleaned_data['department']
                day_letters = form.cleaned_data['day_letters']
                day_digits = form.cleaned_data['day_digits']
                month_letters_a = form.cleaned_data['month_letters_a']
                month_letters_g = form.cleaned_data['month_letters_g']
                year_letters = form.cleaned_data['year_letters']
                year_digits = form.cleaned_data['year_digits']
                notary_office = form.cleaned_data['notary_office']
                name_notary_office = form.cleaned_data['name_notary_office']
                buyer_name = form.cleaned_data['buyer_name']
                number_of_pages = form.cleaned_data['number_of_pages']
                number_form_letters = form.cleaned_data['number_form_letters']
                number_form_digits = form.cleaned_data['number_form_digits']
                amount_to_pay_letters = form.cleaned_data['amount_to_pay_letters']
                amount_to_pay_digits = form.cleaned_data['amount_to_pay_digits']
                she = form.cleaned_data['she']

                if notary_office == 'notario':
                    she = 'el'
                else:
                    she = 'la'

                BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                document_iva_sales_testimony_path = os.path.join(BASE_DIRECTORY, 'documents',
                                                                 'Testimonio_Compraventa_IVA.docx')

                replacements = {
                    "(numero_escritura_letras)": deed_numbers_letters,
                    "(numero_escritura_numeros)": deed_numbers_digits,
                    "(municipio)": municipality,
                    "(departamento)": department,
                    "(dia_letras)": day_letters,
                    "(dia_numeros)": day_digits,
                    "(mes_letras_a)": month_letters_a,
                    "(mes_letras_g)": month_letters_g,
                    "(anio_letras)": year_letters,
                    "(anio_numeros)": year_digits,
                    "(notario_notaria)": notary_office,
                    "(nombre_notario_notaria)": name_notary_office,
                    "(nombre_comprador)": buyer_name,
                    "(numero_hojas)": number_of_pages,
                    "(numero_formulario_letras)": number_form_letters,
                    "(numero_formulario_numeros)": number_form_digits,
                    "(cantidad_pagar_letras)": amount_to_pay_letters,
                    "(cantidad_pagar_numeros)": amount_to_pay_digits,
                    "(el_la)": she
                }

                temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Testimonio_Compraventa_IVA.docx')
                replace_text_in_docx(document_iva_sales_testimony_path, replacements, temp_docx_path)

                request.session['word_generated'] = True
                return redirect('download_iva_sales_testimony_docx')

        else:
            form = IvaSalesTestimonyForm()

        return render(request, 'documents/iva-sales-testimony-document.html', {'form': form})

    finally:
        pythoncom.CoUninitialize()


def generate_municipality_notices_document(request):
    pythoncom.CoInitialize()

    try:
        if request.method == 'POST':
            form = MunicipalityNoticesForm(request.POST)
            if form.is_valid():
                municipality = form.cleaned_data['municipality']
                department = form.cleaned_data['department']
                deed_numbers_letters = form.cleaned_data['deed_numbers_letters']
                deed_numbers_digits = form.cleaned_data['deed_numbers_digits']
                address_property = form.cleaned_data['address_property']
                municipality_property = form.cleaned_data['municipality_property']
                department_property = form.cleaned_data['department_property']
                number_farm_letters = form.cleaned_data['number_farm_letters']
                number_farm_digits = form.cleaned_data['number_farm_digits']
                number_folio_letters = form.cleaned_data['number_folio_letters']
                number_folio_digits = form.cleaned_data['number_folio_digits']
                number_book_letters = form.cleaned_data['number_book_letters']
                number_book_digits = form.cleaned_data['number_book_digits']
                area_state_letters = form.cleaned_data['area_state_letters']
                area_state_digits = form.cleaned_data['area_state_digits']
                value_property_letters = form.cleaned_data['value_property_letters']
                value_property_digits = form.cleaned_data['value_property_digits']
                day_letters = form.cleaned_data['day_letters']
                day_digits = form.cleaned_data['day_digits']
                month_letters_a = form.cleaned_data['month_letters_a']
                month_letters_g = form.cleaned_data['month_letters_g']
                year_letters = form.cleaned_data['year_letters']
                year_digits = form.cleaned_data['year_digits']
                notary_office = form.cleaned_data['notary_office']
                lawyer_office = form.cleaned_data['lawyer_office']
                name_notary_office = form.cleaned_data['name_notary_office']
                collegiate_number = form.cleaned_data['collegiate_number']
                address_receive_notifications = form.cleaned_data['address_receive_notifications']
                municipality_receive_notifications = form.cleaned_data['municipality_receive_notifications']
                department_receive_notifications = form.cleaned_data['department_receive_notifications']
                email_notary_office = form.cleaned_data['email_notary_office']
                seller_name = form.cleaned_data['seller_name']
                seller_dpi_number = form.cleaned_data['seller_dpi_number']
                seller_nit_number = form.cleaned_data['seller_nit_number']
                seller_home_address = form.cleaned_data['seller_home_address']
                seller_home_municipality = form.cleaned_data['seller_home_municipality']
                seller_home_department = form.cleaned_data['seller_home_department']
                seller_registration = form.cleaned_data[
                                          'seller_registration'] or 'No tiene, abrir conforme lo establece el Decreto 15-98 del Congreso de la República.'
                buyer_name = form.cleaned_data['buyer_name']
                buyer_dpi_number = form.cleaned_data['buyer_dpi_number']
                buyer_nit_number = form.cleaned_data['buyer_nit_number']
                buyer_home_address = form.cleaned_data['buyer_home_address']
                buyer_home_municipality = form.cleaned_data['buyer_home_municipality']
                buyer_home_department = form.cleaned_data['buyer_home_department']
                buyer_registration = form.cleaned_data[
                                         'buyer_registration'] or 'No tiene, abrir conforme lo establece el Decreto 15-98 del Congreso de la República.'
                she = form.cleaned_data['she']

                if notary_office == 'notario':
                    she = 'el'
                    notary_office_ms = 'Notario'
                else:
                    she = 'la'
                    notary_office_ms = 'Notaria'

                if lawyer_office == 'abogado':
                    lawyer_office_ms = 'Abogado'
                else:
                    lawyer_office_ms = 'Abogada'

                BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                document_iva_sales_testimony_path = os.path.join(BASE_DIRECTORY, 'documents', 'Avisos_Muni_Dicabi.docx')

                replacements = {
                    "(municipalidad)": municipality,
                    "(departamento)": department,
                    "(numero_escritura_letras)": deed_numbers_letters,
                    "(numero_escritura_digitos)": deed_numbers_digits,
                    "(direccion_inmueble)": address_property,
                    "(municipio_inmueble)": municipality_property,
                    "(departamento_inmueble)": department_property,
                    "(numero_finca_letras)": number_farm_letters,
                    "(numero_finca_digitos)": number_farm_digits,
                    "(numero_folio_letras)": number_folio_letters,
                    "(numero_folio_digitos)": number_folio_digits,
                    "(numero_libro_letras)": number_book_letters,
                    "(numero_libro_digitos)": number_book_digits,
                    "(area_finca_letras)": area_state_letters,
                    "(area_finca_digitos)": area_state_digits,
                    "(valor_inmueble_letras)": value_property_letters,
                    "(valor_inmueble_digitos)": value_property_digits,
                    "(dia_letras)": day_letters,
                    "(dia)": day_digits,
                    "(mes_letras_a)": month_letters_a,
                    "(mes_letras_g)": month_letters_g,
                    "(anio_letras)": year_letters,
                    "(anio)": year_digits,
                    "(notario_notaria_m)": notary_office,
                    "(notario_notaria_ms)": notary_office_ms,
                    "(abogado_abogada_ms)": lawyer_office_ms,
                    "(nombre_notario_notaria)": name_notary_office,
                    "(numero_colegiado)": collegiate_number,
                    "(direccion_recibir_notificaciones)": address_receive_notifications,
                    "(municipio_recibir_notificaciones)": municipality_receive_notifications,
                    "(departamento_recibir_notificaciones)": department_receive_notifications,
                    "(correo_notario_notaria)": email_notary_office,
                    "(nombre_vendedor)": seller_name,
                    "(numero_dpi_vendedor)": seller_dpi_number,
                    "(numero_nit_vendedor)": seller_nit_number,
                    "(direccion_domiciliar_vendedor)": seller_home_address,
                    "(municipio_domiciliar_vendedor)": seller_home_municipality,
                    "(departamento_domiciliar_vendedor)": seller_home_department,
                    "(matricula_vendedor)": seller_registration,
                    "(nombre_comprador)": buyer_name,
                    "(numero_dpi_comprador)": buyer_dpi_number,
                    "(numero_nit_comprador)": buyer_nit_number,
                    "(direccion_domiciliar_comprador)": buyer_home_address,
                    "(municipio_domiciliar_comprador)": buyer_home_municipality,
                    "(departamento_domiciliar_comprador)": buyer_home_department,
                    "(matricula_comprador)": buyer_registration,
                    "(el_la)": she
                }

                temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Avisos_Muni_Dicabi.docx')
                replace_text_in_docx(document_iva_sales_testimony_path, replacements, temp_docx_path)

                request.session['word_generated'] = True
                return redirect('download_municipality_notices_docx')

        else:
            form = MunicipalityNoticesForm()

        return render(request, 'documents/municipality_notices-document.html', {'form': form})

    finally:
        pythoncom.CoUninitialize()


def generate_buy_sell_document(request):
    pythoncom.CoInitialize()

    try:
        if request.method == 'POST':
            form = BuySellForm(request.POST)
            if form.is_valid():
                deed_number_letters = form.cleaned_data['deed_number_letters']
                deed_number_digits = form.cleaned_data['deed_number_digits']
                municipality = form.cleaned_data['municipality']
                day_letters = form.cleaned_data['day_letters']
                month_letters = form.cleaned_data['month_letters']
                year_letters = form.cleaned_data['year_letters']
                notary_name = form.cleaned_data['notary_name']
                notary_office = form.cleaned_data['notary_office']
                seller_name = form.cleaned_data['seller_name']
                seller_age_letters = form.cleaned_data['seller_age_letters']
                seller_civil_status = form.cleaned_data['seller_civil_status']
                seller_nationality = form.cleaned_data['seller_nationality']
                seller_profession = form.cleaned_data['seller_profession']
                seller_department = form.cleaned_data['seller_department']
                seller_dpi_number_letters = form.cleaned_data['seller_dpi_number_letters']
                seller_dpi_number_digits = form.cleaned_data['seller_dpi_number_digits']
                buyer_name = form.cleaned_data['buyer_name']
                buyer_age_letters = form.cleaned_data['buyer_age_letters']
                buyer_civil_status = form.cleaned_data['buyer_civil_status']
                buyer_nationality = form.cleaned_data['buyer_nationality']
                buyer_profession = form.cleaned_data['buyer_profession']
                buyer_department = form.cleaned_data['buyer_department']
                buyer_dpi_number_letters = form.cleaned_data['buyer_dpi_number_letters']
                buyer_dpi_number_digits = form.cleaned_data['buyer_dpi_number_digits']
                farm_number_letters = form.cleaned_data['farm_number_letters']
                farm_number_digits = form.cleaned_data['farm_number_digits']
                folio_number_letters = form.cleaned_data['folio_number_letters']
                folio_number_digits = form.cleaned_data['folio_number_digits']
                book_number_letters = form.cleaned_data['book_number_letters']
                book_number_digits = form.cleaned_data['book_number_digits']
                property_department = form.cleaned_data['property_department']
                municipality_address = form.cleaned_data['municipality_address']
                property_address = form.cleaned_data['property_address']
                property_type = form.cleaned_data['property_type']
                sale_value_letters = form.cleaned_data['sale_value_letters']
                sale_value_digits = form.cleaned_data['sale_value_digits']
                she = form.cleaned_data['she']

                if notary_office == 'notario':
                    she = 'el'
                else:
                    she = 'la'

                BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                document_iva_sales_testimony_path = os.path.join(BASE_DIRECTORY, 'documents', 'Compraventa_Simple.docx')

                replacements = {
                    "(numero_escritura_letras)": deed_number_letters,
                    "(numero_escritura_digitos)": deed_number_digits,
                    "(municipio)": municipality,
                    "(dia_letras)": day_letters,
                    "(mes_letras)": month_letters,
                    "(anio_letras)": year_letters,
                    "(nombre_notario_notaria)": notary_name,
                    "(notario_notaria)": notary_office,
                    "(nombre_vendedor)": seller_name,
                    "(edad_letras_vendedor)": seller_age_letters,
                    "(estado_civil_vendedor)": seller_civil_status,
                    "(nacionalidad_vendedor)": seller_nationality,
                    "(profesion_vendedor)": seller_profession,
                    "(departamento_vendedor)": seller_department,
                    "(numero_dpi_letras_vendedor)": seller_dpi_number_letters,
                    "(numero_dpi_digitos_vededor)": seller_dpi_number_digits,
                    "(nombre_comprador)": buyer_name,
                    "(edad_letras_comprador)": buyer_age_letters,
                    "(estado_civil_comprador)": buyer_civil_status,
                    "(nacionalidad_comprador)": buyer_nationality,
                    "(profesion_comprador)": buyer_profession,
                    "(departamento_comprador)": buyer_department,
                    "(numero_dpi_letras_comprador)": buyer_dpi_number_letters,
                    "(numero_dpi_digitos_comprador)": buyer_dpi_number_digits,
                    "(numero_finca_letras)": farm_number_letters,
                    "(numero_finca_digitos)": farm_number_digits,
                    "(numero_folio_letras)": folio_number_letters,
                    "(numero_folio_digitos)": folio_number_digits,
                    "(numero_libro_letras)": book_number_letters,
                    "(numero_libro_digitos)": book_number_digits,
                    "(departamento_inmueble)": property_department,
                    "(municipio_inmueble)": municipality_address,
                    "(direccion_inmueble)": property_address,
                    "(tipo_finca)": property_type,
                    "(valor_venta_letras)": sale_value_letters,
                    "(valor_venta_digitos)": sale_value_digits,
                    "(el_la)": she
                }

                temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Compraventa_Simple.docx')
                replace_text_in_docx(document_iva_sales_testimony_path, replacements, temp_docx_path)

                request.session['word_generated'] = True
                return redirect('download_buy_sell_docx')

        else:
            form = BuySellForm()

        return render(request, 'documents/buy-sell-document.html', {'form': form})

    finally:
        pythoncom.CoUninitialize()


# Function save file to base64 for database
def save_file_in_base64(path, model_class):
    # Abrir el archivo en binario
    with open(path, 'rb') as file:
        # Convertir a base64
        encoded_file = base64.b64encode(file.read()).decode('utf-8')
        # Guardar en la base de datos
        model_class.objects.create(file_base64=encoded_file)


# Views downloads files Word
def download_special_testimony_docx(request):
    if request.session.get('word_generated', False):
        del request.session['word_generated']

        BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Testimonio_Especial.docx')

        if os.path.exists(temp_docx_path):
            save_file_in_base64(temp_docx_path, SpecialTestimony)

            with open(temp_docx_path, 'rb') as docx_file:
                response = HttpResponse(docx_file.read(),
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="Testimonio_Especial.docx"'

            os.remove(temp_docx_path)

            return response
        else:
            return HttpResponse("El archivo Word no se encontró.", status=404)

    return redirect(reverse('generate_special_testimony_document'))


def download_iva_sales_testimony_docx(request):
    if request.session.get('word_generated', False):
        del request.session['word_generated']

        BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Testimonio_Compraventa_IVA.docx')

        if os.path.exists(temp_docx_path):
            save_file_in_base64(temp_docx_path, IvaSalesTestimony)

            with open(temp_docx_path, 'rb') as docx_file:
                response = HttpResponse(docx_file.read(),
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="Testimonio_Compraventa_IVA.docx"'

            os.remove(temp_docx_path)

            return response
        else:
            return HttpResponse("El archivo Word no se encontró.", status=404)

    return redirect(reverse('generate_iva_sales_testimony_document'))


def download_municipality_notices_docx(request):
    if request.session.get('word_generated', False):
        del request.session['word_generated']

        BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Avisos_Muni_Dicabi.docx')

        if os.path.exists(temp_docx_path):
            save_file_in_base64(temp_docx_path, MunicipalityNotices)

            with open(temp_docx_path, 'rb') as docx_file:
                response = HttpResponse(docx_file.read(),
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="Avisos_Muni_Dicabi.docx"'

            os.remove(temp_docx_path)

            return response
        else:
            return HttpResponse("El archivo Word no se encontró.", status=404)

    return redirect(reverse('generate_municipality_notices_document'))


def download_buy_sell_docx(request):
    if request.session.get('word_generated', False):
        del request.session['word_generated']

        BASE_DIRECTORY = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        temp_docx_path = os.path.join(BASE_DIRECTORY, 'documents', 'Temp_Compraventa_Simple.docx')

        if os.path.exists(temp_docx_path):
            save_file_in_base64(temp_docx_path, BuySell)

            with open(temp_docx_path, 'rb') as docx_file:
                response = HttpResponse(docx_file.read(),
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="Compraventa_Simple.docx"'

            os.remove(temp_docx_path)

            return response
        else:
            return HttpResponse("El archivo Word no se encontró.", status=404)

    return redirect(reverse('generate_buy_sell_document'))
